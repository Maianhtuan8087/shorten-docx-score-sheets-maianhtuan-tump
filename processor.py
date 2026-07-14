#!/usr/bin/env python3
"""Convert DOCX answer keys into compact grading sheets with a QC report."""

from __future__ import annotations

import argparse
import json
import re
import sys
from copy import deepcopy
from dataclasses import dataclass
from decimal import Decimal, InvalidOperation
from pathlib import Path
from typing import Any

from docx import Document
from docx.oxml.ns import qn
from docx.table import _Cell, Table


ELLIPSIS_RE = re.compile(r"(?<!\.)\.{3,}(?!\.)|…")
SCORE_RE = re.compile(r"^\s*0\s*,\s*25\s*(?:đ|điểm)?\s*$", re.IGNORECASE)
QUESTION_RE = re.compile(r"^\s*Câu\s*(\d+)\s*[.:]?\s*(.*)$", re.IGNORECASE)
POINT_RE = re.compile(r"\(\s*(\d+(?:[,.]\d+)?)\s*điểm\s*\)", re.IGNORECASE)
HEADERS = {
    "nội dung",
    "điểm",
    "thang điểm",
    "nội dung trả lời",
    "nội dung đáp án",
    "câu số",
    "điểm chấm",
}
_SCRIPT_PATH = globals().get("__file__")
_SCRIPT_DIR = Path(_SCRIPT_PATH).resolve().parent if _SCRIPT_PATH else Path.cwd()
_PROJECT_TEMPLATE = _SCRIPT_DIR / "assets" / "grading-sheet-template.docx"
_SKILL_TEMPLATE = _SCRIPT_DIR.parent / "assets" / "grading-sheet-template.docx"
DEFAULT_TEMPLATE = _PROJECT_TEMPLATE if _PROJECT_TEMPLATE.exists() else _SKILL_TEMPLATE


@dataclass
class Question:
    number: int
    text: str
    points: Decimal | None


@dataclass
class RubricItem:
    text: str
    score: Decimal
    table_index: int
    row_index: int


def normalize_space(text: str) -> str:
    """Collapse all Word paragraph/line whitespace to single spaces."""
    return re.sub(r"\s+", " ", text or "").strip()


def tokenize_words_preserve_punctuation(text: str) -> list[str]:
    """Split on whitespace while keeping punctuation attached to each token."""
    return normalize_space(text).split() if normalize_space(text) else []


def count_words(text: str) -> int:
    return len(tokenize_words_preserve_punctuation(text))


def is_score_cell(text: str) -> bool:
    return bool(SCORE_RE.fullmatch(normalize_space(text)))


def is_header_text(text: str) -> bool:
    return normalize_space(text).casefold() in HEADERS


def has_ellipsis(text: str) -> bool:
    return bool(ELLIPSIS_RE.search(text or ""))


def is_already_shortened(text: str) -> bool:
    """Recognize a compact intentional ellipsis, not any natural ellipsis."""
    compact = normalize_space(text)
    matches = list(ELLIPSIS_RE.finditer(compact))
    if len(matches) != 1:
        return False
    match = matches[0]
    before = tokenize_words_preserve_punctuation(compact[: match.start()])
    after = tokenize_words_preserve_punctuation(compact[match.end() :])
    return 4 <= len(before) <= 7 and 2 <= len(after) <= 5 and len(before) + len(after) <= 12


def shorten_text(text: str, first_n: int = 5, last_n: int = 3) -> tuple[str, str]:
    """Return (text, status): shortened/already_shortened/too_short/empty."""
    compact = normalize_space(text)
    if not compact:
        return "", "empty"
    words = tokenize_words_preserve_punctuation(compact)
    if len(words) <= first_n + last_n:
        return compact, "too_short"
    if is_already_shortened(compact):
        return compact, "already_shortened"
    return " ".join(words[:first_n] + ["..."] + words[-last_n:]), "shortened"


def replace_cell_text_preserve_basic_format(cell: _Cell, new_text: str) -> None:
    """Replace cell text while retaining its main paragraph/run formatting."""
    paragraph = cell.paragraphs[0] if cell.paragraphs else cell.add_paragraph()
    run = paragraph.runs[0] if paragraph.runs else None
    snapshot: dict[str, Any] = {
        "alignment": paragraph.alignment,
        "style": paragraph.style,
        "vertical_alignment": cell.vertical_alignment,
    }
    if run is not None:
        snapshot.update(
            {
                "bold": run.bold,
                "italic": run.italic,
                "underline": run.underline,
                "name": run.font.name,
                "size": run.font.size,
                "color": run.font.color.rgb,
            }
        )
    cell.text = new_text
    paragraph = cell.paragraphs[0]
    paragraph.alignment = snapshot["alignment"]
    if snapshot["style"] is not None:
        paragraph.style = snapshot["style"]
    cell.vertical_alignment = snapshot["vertical_alignment"]
    run = paragraph.runs[0] if paragraph.runs else paragraph.add_run()
    for key in ("bold", "italic", "underline"):
        if key in snapshot:
            setattr(run, key, snapshot[key])
    if snapshot.get("name"):
        run.font.name = snapshot["name"]
    if snapshot.get("size"):
        run.font.size = snapshot["size"]
    if snapshot.get("color"):
        run.font.color.rgb = snapshot["color"]


def _decimal(text: str) -> Decimal:
    try:
        return Decimal(text.replace(",", ".").strip())
    except InvalidOperation as exc:
        raise ValueError(f"Invalid score: {text!r}") from exc


def _score_value(text: str) -> Decimal:
    match = re.search(r"\d+(?:[,.]\d+)?", normalize_space(text))
    if not match:
        raise ValueError(f"Missing numeric score in {text!r}")
    return _decimal(match.group(0))


def _parse_questions(document) -> list[Question]:
    questions: list[Question] = []
    for paragraph in document.paragraphs:
        text = normalize_space(paragraph.text)
        match = QUESTION_RE.match(text)
        if not match:
            continue
        number = int(match.group(1))
        body = match.group(2)
        point_match = POINT_RE.search(body)
        points = _decimal(point_match.group(1)) if point_match else None
        body = normalize_space(POINT_RE.sub("", body))
        questions.append(Question(number, f"{number}. {body}".strip(), points))
    return questions


def _find_scoring_cell_index(row) -> int | None:
    for index, cell in enumerate(row.cells):
        if is_score_cell(cell.text):
            return index
    return None


def _content_cell_index(row, score_index: int) -> int | None:
    if score_index > 0:
        candidate = row.cells[score_index - 1]
        if normalize_space(candidate.text) and not is_header_text(candidate.text):
            return score_index - 1
    for index, cell in enumerate(row.cells):
        text = normalize_space(cell.text)
        if index != score_index and text and not is_header_text(text) and not is_score_cell(text):
            return index
    return None


def _extract_rubrics(document) -> list[list[RubricItem]]:
    groups: list[list[RubricItem]] = []
    for table_index, table in enumerate(document.tables):
        items: list[RubricItem] = []
        for row_index, row in enumerate(table.rows):
            score_index = _find_scoring_cell_index(row)
            if score_index is None:
                continue
            content_index = _content_cell_index(row, score_index)
            if content_index is None:
                continue
            items.append(
                RubricItem(
                    normalize_space(row.cells[content_index].text),
                    _score_value(row.cells[score_index].text),
                    table_index,
                    row_index,
                )
            )
        if items:
            groups.append(items)
    return groups


def _is_question_table(table: Table) -> bool:
    if len(table.columns) != 4 or not table.rows:
        return False
    headers = [normalize_space(cell.text).casefold() for cell in table.rows[0].cells]
    return "câu số" in headers and "điểm chấm" in headers


def _table_text(table: Table) -> str:
    return "\n".join(cell.text for row in table.rows for cell in row.cells)


def _raw_cell(table: Table, row_index: int, column_index: int) -> _Cell:
    tc = table.rows[row_index]._tr.tc_lst[column_index]
    return _Cell(tc, table)


def _remove_vertical_merges(table: Table, start_row: int, end_row: int) -> None:
    for row_index in range(start_row, end_row + 1):
        tc = table.rows[row_index]._tr.tc_lst[0]
        tc_pr = tc.get_or_add_tcPr()
        vmerge = tc_pr.find(qn("w:vMerge"))
        if vmerge is not None:
            tc_pr.remove(vmerge)


def _resize_question_table(table: Table, answer_count: int) -> None:
    if len(table.rows) < 3:
        raise ValueError("Template question table must have header, answer, and total rows")
    _remove_vertical_merges(table, 1, len(table.rows) - 2)
    target_rows = answer_count + 2
    while len(table.rows) > target_rows:
        row = table.rows[-2]._tr
        row.getparent().remove(row)
    while len(table.rows) < target_rows:
        clone = deepcopy(table.rows[-2]._tr)
        table.rows[-1]._tr.addprevious(clone)
    _remove_vertical_merges(table, 1, len(table.rows) - 2)


def _set_question_table(
    table: Table,
    question: Question,
    items: list[RubricItem],
    report: dict[str, Any],
) -> None:
    _resize_question_table(table, len(items))
    headers = ["Câu số", "Nội dung", "Thang điểm", "Điểm chấm"]
    for col, value in enumerate(headers):
        replace_cell_text_preserve_basic_format(_raw_cell(table, 0, col), value)

    for offset, item in enumerate(items, start=1):
        new_text, status = shorten_text(item.text)
        if status == "shortened":
            report["cells_shortened"] += 1
            if has_ellipsis(item.text):
                report["natural_ellipsis_processed"] += 1
        elif status == "already_shortened":
            report["already_shortened_skipped"] += 1
        elif status == "too_short":
            report["too_short_skipped"] += 1
        else:
            report["empty_skipped"] += 1
        replace_cell_text_preserve_basic_format(_raw_cell(table, offset, 0), "")
        replace_cell_text_preserve_basic_format(_raw_cell(table, offset, 1), new_text)
        replace_cell_text_preserve_basic_format(_raw_cell(table, offset, 2), "0,25 điểm")
        replace_cell_text_preserve_basic_format(_raw_cell(table, offset, 3), "")

    if items:
        merged_question_cell = table.cell(1, 0).merge(table.cell(len(items), 0))
        # Set text after merging so the empty continuation cells do not leave
        # seven extra paragraphs that expand the table and change pagination.
        replace_cell_text_preserve_basic_format(merged_question_cell, question.text)

    total = question.points if question.points is not None else sum((x.score for x in items), Decimal("0"))
    total_text = f"{int(total)} điểm" if total == int(total) else f"{str(total).replace('.', ',')} điểm"
    total_row = len(table.rows) - 1
    replace_cell_text_preserve_basic_format(table.cell(total_row, 0), f"Tổng điểm câu {question.number}")
    replace_cell_text_preserve_basic_format(table.cell(total_row, 2), total_text)
    replace_cell_text_preserve_basic_format(table.cell(total_row, 3), "")


def _resize_template_question_tables(document, count: int) -> list[Table]:
    question_tables = [table for table in document.tables if _is_question_table(table)]
    if not question_tables:
        raise ValueError("Template has no 4-column grading table")
    signature_tables = [
        table for table in document.tables if len(table.columns) == 2 and "Điểm toàn bài" in _table_text(table)
    ]
    if not signature_tables:
        raise ValueError("Template has no signature/whole-score table")
    signature = signature_tables[-1]
    prototype = question_tables[-1]._tbl
    while len(question_tables) < count:
        signature._tbl.addprevious(deepcopy(prototype))
        question_tables = [table for table in document.tables if _is_question_table(table)]
    while len(question_tables) > count:
        table = question_tables.pop()
        table._element.getparent().remove(table._element)
    return [table for table in document.tables if _is_question_table(table)]


def _replace_paragraph_text_preserve_format(paragraph, text: str) -> None:
    if paragraph.runs:
        first = paragraph.runs[0]
        snapshot = (first.bold, first.italic, first.underline, first.font.name, first.font.size)
    else:
        snapshot = (None, None, None, None, None)
    paragraph.text = text
    run = paragraph.runs[0] if paragraph.runs else paragraph.add_run()
    run.bold, run.italic, run.underline = snapshot[:3]
    if snapshot[3]:
        run.font.name = snapshot[3]
    if snapshot[4]:
        run.font.size = snapshot[4]


def _update_template_metadata(source, output, questions: list[Question]) -> None:
    source_texts = [normalize_space(p.text) for p in source.paragraphs if normalize_space(p.text)]
    source_title = next((x for x in source_texts if "ĐÁP ÁN" in x.upper()), "")
    source_subject = next((x for x in source_texts if x.casefold().startswith("môn thi")), "")
    title = re.sub("ĐÁP ÁN", "PHIẾU CHẤM", source_title, flags=re.IGNORECASE) if source_title else ""
    for paragraph in output.paragraphs:
        current = normalize_space(paragraph.text)
        if title and "PHIẾU CHẤM" in current.upper():
            _replace_paragraph_text_preserve_format(paragraph, title)
        elif source_subject and current.casefold().startswith("môn thi"):
            _replace_paragraph_text_preserve_format(paragraph, source_subject)

    if source.tables and output.tables and len(source.tables[0].columns) == len(output.tables[0].columns) == 2:
        for column in range(2):
            replace_cell_text_preserve_basic_format(output.tables[0].cell(0, column), source.tables[0].cell(0, column).text)

    year_match = re.search(r"\b(20\d{2})\b", source_title)
    if year_match:
        year = year_match.group(1)
        for table in output.tables:
            if len(table.columns) == 2 and "Điểm toàn bài" in _table_text(table):
                cell = table.cell(0, 0)
                replace_cell_text_preserve_basic_format(cell, re.sub(r"20\d{2}", year, cell.text))


def _select_items(question: Question, items: list[RubricItem], report: dict[str, Any]) -> list[RubricItem]:
    if question.points is None:
        report["warnings"].append(
            f"Câu {question.number}: không tìm thấy tổng điểm trong câu hỏi; giữ toàn bộ {len(items)} dòng 0,25."
        )
        return items
    max_rows = int(question.points / Decimal("0.25"))
    if len(items) > max_rows:
        overflow = items[max_rows:]
        report["overflow_rows_omitted"] += len(overflow)
        report["warnings"].append(
            f"Câu {question.number}: có {len(items)} dòng 0,25 nhưng tổng {question.points} điểm chỉ cho phép {max_rows}; "
            f"đã bỏ {len(overflow)} dòng cuối và ghi nhận trong QC."
        )
    elif len(items) < max_rows:
        report["warnings"].append(
            f"Câu {question.number}: chỉ có {len(items)} dòng 0,25, thấp hơn {max_rows} dòng theo tổng điểm."
        )
    return items[:max_rows]


def process_docx(input_path: str, output_path: str, template_path: str | None = None) -> dict[str, Any]:
    input_file = Path(input_path).resolve()
    output_file = Path(output_path).resolve()
    template_file = Path(template_path).resolve() if template_path else DEFAULT_TEMPLATE
    report: dict[str, Any] = {
        "input_file": str(input_file),
        "output_file": str(output_file),
        "template_file": str(template_file),
        "tables_scanned": 0,
        "questions_found": 0,
        "score_rows_found": 0,
        "score_rows_written": 0,
        "overflow_rows_omitted": 0,
        "cells_shortened": 0,
        "already_shortened_skipped": 0,
        "too_short_skipped": 0,
        "empty_skipped": 0,
        "natural_ellipsis_processed": 0,
        "warnings": [],
        "errors": [],
    }
    try:
        if input_file == output_file:
            raise ValueError("Output must be a new file; input and output paths are identical")
        if not input_file.exists():
            raise FileNotFoundError(input_file)
        if not template_file.exists():
            raise FileNotFoundError(template_file)
        source = Document(str(input_file))
        output = Document(str(template_file))
        questions = _parse_questions(source)
        rubrics = _extract_rubrics(source)
        report["tables_scanned"] = len(source.tables)
        report["questions_found"] = len(questions)
        report["score_rows_found"] = sum(len(group) for group in rubrics)
        if not questions:
            raise ValueError("No question paragraphs matching 'Câu <number>' were found")
        if len(questions) != len(rubrics):
            raise ValueError(f"Question/table mismatch: {len(questions)} questions, {len(rubrics)} scoring tables")
        tables = _resize_template_question_tables(output, len(questions))
        _update_template_metadata(source, output, questions)
        for question, items, table in zip(questions, rubrics, tables):
            selected = _select_items(question, items, report)
            _set_question_table(table, question, selected, report)
            report["score_rows_written"] += len(selected)
        output_file.parent.mkdir(parents=True, exist_ok=True)
        output.save(str(output_file))
    except Exception as exc:  # CLI must return a readable QC report.
        report["errors"].append(f"{type(exc).__name__}: {exc}")
    report_path = output_file.with_suffix(".qc.json")
    report_path.parent.mkdir(parents=True, exist_ok=True)
    report_path.write_text(json.dumps(report, ensure_ascii=False, indent=2), encoding="utf-8")
    report["report_file"] = str(report_path)
    return report


def process_folder(input_folder: str, output_folder: str, template_path: str | None = None) -> list[dict[str, Any]]:
    source_dir = Path(input_folder).resolve()
    target_dir = Path(output_folder).resolve()
    target_dir.mkdir(parents=True, exist_ok=True)
    reports = []
    for input_file in sorted(source_dir.glob("*.docx")):
        if input_file.name.startswith("~$"):
            continue
        output_file = target_dir / f"{input_file.stem}_phieu_cham.docx"
        reports.append(process_docx(str(input_file), str(output_file), template_path))
    (target_dir / "processing_report.json").write_text(
        json.dumps(reports, ensure_ascii=False, indent=2), encoding="utf-8"
    )
    return reports


def _build_parser() -> argparse.ArgumentParser:
    parser = argparse.ArgumentParser(description=__doc__)
    mode = parser.add_mutually_exclusive_group(required=True)
    mode.add_argument("--input", help="One input DOCX")
    mode.add_argument("--input-folder", help="Folder containing DOCX files")
    parser.add_argument("--output", help="Output DOCX for one-file mode")
    parser.add_argument("--output-folder", help="Output folder for batch mode")
    parser.add_argument("--template", help="Approved grading-sheet template DOCX")
    return parser


def main() -> int:
    if hasattr(sys.stdout, "reconfigure"):
        sys.stdout.reconfigure(encoding="utf-8")
    args = _build_parser().parse_args()
    if args.input:
        if not args.output:
            print("--output is required with --input", file=sys.stderr)
            return 2
        result: Any = process_docx(args.input, args.output, args.template)
        code = 1 if result["errors"] else 0
    else:
        if not args.output_folder:
            print("--output-folder is required with --input-folder", file=sys.stderr)
            return 2
        result = process_folder(args.input_folder, args.output_folder, args.template)
        code = 1 if any(item["errors"] for item in result) else 0
    print(json.dumps(result, ensure_ascii=False, indent=2))
    return code


if __name__ == "__main__":
    raise SystemExit(main())
